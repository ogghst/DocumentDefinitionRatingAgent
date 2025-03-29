import os
import pandas as pd
import json
import traceback
from datetime import datetime
from dotenv import load_dotenv
from typing import List, Dict, Optional, TypedDict, Any

# Langchain & Langgraph Imports
from langchain_community.document_loaders import Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI # Use ChatOpenAI configured for DeepSeek endpoint
from langchain.prompts import PromptTemplate
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.exceptions import OutputParserException
from langchain_core.runnables import RunnableLambda, RunnablePassthrough, RunnableConfig
from langchain_core.documents import Document
from langchain.vectorstores.base import VectorStoreRetriever
from langgraph.graph import StateGraph, END

from langchain_community.llms import Ollama
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_ollama import ChatOllama, OllamaEmbeddings

# Pydantic Models for Structure
from pydantic import BaseModel, Field, validator, ValidationError, field_validator
from typing_extensions import Annotated

# --- Environment Variable Loading ---
load_dotenv()

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
DEEPSEEK_API_BASE = os.getenv("DEEPSEEK_API_BASE")
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL")

# --- Basic Validation ---
if not DEEPSEEK_API_KEY:
    raise ValueError("DEEPSEEK_API_KEY environment variable not set.")
if not DEEPSEEK_API_BASE:
    raise ValueError("DEEPSEEK_API_BASE environment variable not set.")
if not OLLAMA_MODEL:
    raise ValueError("OLLAMA_MODEL environment variable not set (e.g., 'phi4:14b').")



# --- Pydantic Models for Data Structure ---

class CheckItem(BaseModel):
    """Represents a single checklist item for a specific phase."""
    id: int = Field(..., description="Unique identifier for the check.")
    name: str = Field(..., description="Short name or title of the check.")
    branch_id: int = Field(..., description="Identifier for the check's branch/topic.")
    branch_name: str = Field(..., description="Name of the check's branch/topic.")
    description: str = Field(..., description="Detailed description of the check to be performed.")
    weight: int = Field(..., description="Weight or importance of the check for this phase.")
    phase: str = Field(..., description="The project phase this check instance belongs to.")

class CheckResult(BaseModel):
    """Represents the analysis result for a single check item based on document review."""
    check_item: Optional[CheckItem] = Field(None, description="The original checklist item being analyzed. Can be None if parsing fails early.")
    is_met: Optional[bool] = Field(None, description="True if the document confirms the check is met, False if it's not met or evidence is missing/unclear.")
    reliability: float = Field(..., ge=0.0, le=100.0, description="Confidence score (0-100) that the 'is_met' field is correct based *only* on the provided document context. High score requires direct, unambiguous evidence.")
    sources: List[str] = Field(default_factory=list, description="Specific sentences or passages from the document context that directly support the 'is_met' conclusion and reliability score. Should be exact quotes.")
    analysis_details: str = Field("", description="Brief LLM reasoning or explanation for the conclusion and reliability score, referencing the context.")
    needs_human_review: bool = Field(False, description="True if reliability score is below 50%, indicating need for manual verification.")

    @field_validator('reliability')
    def check_reliability_range(cls, v):
        if not (0 <= v <= 100):
            # Clamp the value instead of raising an error, as LLMs might slightly exceed bounds
            print(f"Warning: Clamping reliability score {v} to be within [0, 100].")
            return max(0.0, min(100.0, v))
        return v

# --- LangGraph State Definition ---

class GraphState(TypedDict):
    """Represents the state passed between nodes in the LangGraph workflow."""
    # Inputs
    checklist_path: Annotated[str, "Path to the Excel checklist file."]
    document_path: Annotated[str, "Path to the Word quotation document."]
    target_phase: Annotated[str, "The specific project phase to analyze checks for."]

    # Processed data
    checks_for_phase: Annotated[Optional[List[CheckItem]], "List of CheckItem objects relevant to the target_phase."] = None
    retriever: Annotated[Optional[VectorStoreRetriever], "Retriever built from the document chunks."] = None

    # Results accumulation
    # This key ('analysis_results') will hold the output of the mapped analysis node
    analysis_results: Annotated[Optional[List[CheckResult]], "List to accumulate analysis results for each check."] = None
    final_results: Annotated[Optional[List[Dict]], "Final list of results as dictionaries for JSON output."] = None
    error_message: Annotated[Optional[str], "To capture any errors during processing."] = None

# --- LLM, Embeddings, and Parser Initialization ---

try:

    print(f"Using Ollama Embeddings: model='{OLLAMA_MODEL}', base_url='{OLLAMA_BASE_URL}'")

    # Initialize Ollama Embeddings
    embeddings = OllamaEmbeddings(
        model=OLLAMA_MODEL,
        base_url=OLLAMA_BASE_URL
    )
    # Test embedding connection (optional but recommended)
    # embeddings.embed_query("Test connection")
    print("Ollama Embeddings initialized successfully.")

    # Initialize ChatOpenAI client configured for DeepSeek
    # print(f"Using DeepSeek LLM via endpoint: '{DEEPSEEK_API_BASE}'")
    # Make sure DeepSeek uses 'openai_api_key' and 'openai_api_base' style configuration
    #llm = ChatOpenAI(
    #    model="deepseek-chat", # Use the appropriate model name for DeepSeek
    #    # model="deepseek-coder", # Alternative if using coder model
    #    temperature=0.0, # Low temperature for factual analysis
    #    api_key=DEEPSEEK_API_KEY,
    #    base_url=DEEPSEEK_API_BASE,
    #    max_tokens=1024, # Adjust as needed
    #)

    print(f"Using Ollama LLM: model='{OLLAMA_MODEL}', base_url='{OLLAMA_BASE_URL}'")

    llm = ChatOllama(
        model=OLLAMA_MODEL,
        base_url=OLLAMA_BASE_URL,
        max_tokens=1024,
        temperature=0.0,
        max_predict=1024,
    )



    # Test LLM connection (optional)
    # llm.invoke("Hello!")
    print("DeepSeek LLM client initialized successfully.")

except Exception as e:
    print(f"ERROR initializing LangChain components: {e}")
    print("Please ensure Ollama is running, the embedding model is pulled,")
    print("and DeepSeek API key/base URL are correct in the .env file.")
    raise

# Initialize the output parser for the analysis result
analysis_parser = PydanticOutputParser(pydantic_object=CheckResult)

# --- Node Functions ---

def load_and_filter_checklist(state: GraphState) -> GraphState:
    """Loads checklist from Excel and filters checks for the target phase."""
    print(f"\n--- Node: load_and_filter_checklist ---")
    checklist_path = state['checklist_path']
    target_phase = state['target_phase']
    print(f"Loading Checklist: '{checklist_path}' for Phase: '{target_phase}'")
    try:
        if not os.path.exists(checklist_path):
            raise FileNotFoundError(f"Checklist file not found at: {checklist_path}")

        df = pd.read_excel(checklist_path)
        # Standardize column names (lowercase, replace spaces) for robustness
        df.columns = df.columns.str.lower().str.replace(' ', '_')

        required_cols = ['id', 'name', 'branchid', 'branchname', 'chk_description', 'weight', 'phase']
        # Check based on standardized names
        if not all(col in df.columns for col in required_cols):
            missing = [col for col in required_cols if col not in df.columns]
            raise ValueError(f"Checklist missing required columns (standardized names): {missing}")

        # Filter by phase (case-insensitive, strip whitespace)
        target_phase_lower = target_phase.strip().lower()
        filtered_df = df[df['phase'].astype(str).str.strip().str.lower() == target_phase_lower]

        checks = []
        for _, row in filtered_df.iterrows():
             try:
                 checks.append(CheckItem(
                     id=int(row['id']), # Ensure correct types
                     name=str(row['name']),
                     branch_id=int(row['branchid']),
                     branch_name=str(row['branchname']),
                     description=str(row['chk_description']),
                     weight=int(row['weight']),
                     phase=str(row['phase'])
                 ))
             except (ValueError, TypeError) as ve:
                 print(f"Warning: Skipping row due to type conversion error: {ve} - Row data: {row.to_dict()}")
                 continue # Skip rows with bad data

        print(f"Found {len(checks)} checks for phase '{target_phase}'.")
        state["checks_for_phase"] = checks
        state["error_message"] = None # Clear previous errors if successful
    except Exception as e:
        print(f"ERROR loading checklist: {e}")
        traceback.print_exc()
        state["error_message"] = f"Failed to load/filter checklist: {e}"
        state["checks_for_phase"] = [] # Ensure it's an empty list on error
    return state

def load_index_document(state: GraphState) -> GraphState:
    """Loads the Word document, chunks it, creates embeddings, and sets up the retriever."""
    print(f"\n--- Node: load_index_document ---")
    if state.get("error_message"):
        print("Skipping due to previous error.")
        return state # Skip if previous step failed

    document_path = state['document_path']
    print(f"Loading & Indexing Document: '{document_path}'")
    try:
        if not os.path.exists(document_path):
            raise FileNotFoundError(f"Document file not found at: {document_path}")

        loader = Docx2txtLoader(document_path)
        raw_docs = loader.load()
        if not raw_docs or not raw_docs[0].page_content.strip():
             raise ValueError("Document is empty or could not be loaded/parsed.")

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,  # Size of chunks
            chunk_overlap=150, # Overlap between chunks
            length_function=len,
            add_start_index=True, # Helps in locating sources if needed later
        )
        documents = text_splitter.split_documents(raw_docs)
        if not documents:
             raise ValueError("Document processed into zero chunks. Check content and splitter settings.")

        print(f"Split document into {len(documents)} chunks.")

        # Create vector store using Ollama embeddings and FAISS
        print("Creating vector store with Ollama embeddings...")
        vectorstore = FAISS.from_documents(documents, embeddings)
        state["retriever"] = vectorstore.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 4} # Retrieve top 4 most relevant chunks
        )
        print("Document indexed successfully. Retriever is ready.")
        state["error_message"] = None # Clear previous errors
    except Exception as e:
        print(f"ERROR loading/indexing document: {e}")
        traceback.print_exc()
        state["error_message"] = f"Failed to load or index document: {e}"
        state["retriever"] = None # Ensure retriever is None on error
    return state

def analyze_check_rag(check_item: CheckItem, config: RunnableConfig) -> CheckResult:
    """
    Analyzes a single check item using RAG against the indexed document.
    Designed to be called via LangGraph's map functionality.
    """
    print(f"\n--- Analyzing Check ID: {check_item.id} ({check_item.name}) ---")
    # Access the shared state (including the retriever) from the config
    state: GraphState = config['configurable']
    retriever = state.get('retriever')

    # --- Pre-computation Checks ---
    if not retriever:
        print(f"ERROR: Retriever not available for check ID {check_item.id}. Skipping analysis.")
        return CheckResult(
            check_item=check_item, is_met=None, reliability=0.0, sources=[],
            analysis_details="Analysis skipped: Document retriever not initialized.",
            needs_human_review=True
        )

    # --- Define Prompt Template ---
    prompt_template = """
    You are a meticulous Project Compliance Analyst. Your task is to analyze the provided Document Context to determine if a specific Checklist Item is met for a given project phase.

    **Checklist Item Details:**
    - ID: {check_id}
    - Name: {check_name}
    - Description: {check_description}
    - Phase: {check_phase}

    **Document Context:**
    --- START CONTEXT ---
    {context}
    --- END CONTEXT ---
    
    {additional_info}

    **Analysis Instructions:**
    1.  **Understand the Goal:** Read the Checklist Item Description carefully. What specific condition needs to be confirmed?
    2.  **Examine Context:** Search the Document Context for explicit statements or strong evidence directly related to the checklist item's condition.
    3.  **Determine Status (`is_met`):**
        *   If the context clearly and unambiguously confirms the condition is met, set `is_met` to `true`.
        *   If the context clearly confirms the condition is *not* met, or if the context lacks any relevant information or is too vague to make a determination, set `is_met` to `false`.
    4.  **Assess Reliability (`reliability`):** Provide a confidence score (0-100) based *only* on the provided context:
        *   90-100: Direct, explicit confirmation/denial in the context. No ambiguity.
        *   70-89: Strong evidence suggesting confirmation/denial, but requires minimal interpretation.
        *   50-69: Context provides related information, but it's indirect or requires significant interpretation. Plausible but not certain.
        *   0-49: Context is irrelevant, contradictory, very vague, or completely missing information about the check item.
    5.  **Extract Sources (`sources`):** Quote the *exact* sentences or short passages (max 2-3 relevant sentences per source) from the Document Context that are the primary evidence for your `is_met` decision and `reliability` score. If no specific sentences provide direct evidence, provide an empty list `[]`.
    6.  **Explain Reasoning (`analysis_details`):** Briefly explain *why* you reached the `is_met` conclusion and `reliability` score, referencing the evidence (or lack thereof) in the context.

    **IMPORTANT: Do NOT include the check_item field in your response. Only provide is_met, reliability, sources, and analysis_details.**

    **Output Format:**
    {format_instructions}
    """

    # --- Define Interactive Question Generation Prompt ---
    question_prompt_template = """
    You are a meticulous Project Compliance Analyst. You're reviewing a checklist item but need additional information to make a confident determination.

    **Checklist Item Details:**
    - ID: {check_id}
    - Name: {check_name}
    - Description: {check_description}
    - Phase: {check_phase}

    **Current Analysis Status:**
    - Current reliability score: {current_reliability}%
    - Current determination: {is_met_status}
    - Current reasoning: {analysis_details}

    **Document Evidence:**
    {sources_summary}

    Based on the above information, identify 1-3 specific questions that would help clarify whether this checklist item is met or not. Focus on questions that:
    1. Target the precise information gaps in the document
    2. Would significantly increase the reliability of your assessment
    3. Are directly relevant to determining if the checklist item is satisfied

    Format your response as a numbered list of questions only. Do not provide any additional text or explanations.
    """

    question_prompt = PromptTemplate(
        template=question_prompt_template,
        input_variables=["check_id", "check_name", "check_description", "check_phase", 
                        "current_reliability", "is_met_status", "analysis_details", 
                        "sources_summary"]
    )

    prompt = PromptTemplate(
        template=prompt_template,
        input_variables=["check_id", "check_name", "check_description", "check_phase", "context", "additional_info"],
        partial_variables={"format_instructions": analysis_parser.get_format_instructions()}
    )

    # --- Define RAG Chain ---
    def format_docs(docs: List[Document]) -> str:
        """Helper to join document contents."""
        return "\n\n".join(f"Source Chunk {i+1}:\n{doc.page_content}" for i, doc in enumerate(docs))

    rag_chain = (
        RunnablePassthrough.assign(
            # Retrieve documents based on the check description
            context=(lambda x: x["check_description"]) | retriever | format_docs
        )
        | prompt
        | llm
        | analysis_parser # Parse the LLM output into the CheckResult Pydantic model
    )

    # --- Invoke Chain and Handle Results ---
    try:
        # Prepare input dictionary for the chain
        additional_info = ""  # Start with no additional info
        max_attempts = 3      # Maximum number of interactive attempts
        attempt = 0
        
        while attempt < max_attempts:
            # Prepare input dictionary for the chain
            chain_input = {
                "check_id": check_item.id,
                "check_name": check_item.name,
                "check_description": check_item.description,
                "check_phase": check_item.phase,
                "additional_info": additional_info
            }
            
            # Invoke the chain
            result = rag_chain.invoke(chain_input, config=config)
            
            # Make sure to add the check_item AFTER parsing the LLM response
            result.check_item = check_item
            
            # Check if reliability is below threshold
            if result.reliability < 50 and attempt < max_attempts - 1:
                print(f"-> Check ID {check_item.id}: Reliability {result.reliability:.1f}% < 50%. Gathering more information...")
                
                # Generate questions for the user
                sources_summary = "No relevant sources found in the document." if not result.sources else "\n".join(result.sources)
                is_met_status = "Undetermined" if result.is_met is None else ("Met" if result.is_met else "Not Met")
                
                question_input = {
                    "check_id": check_item.id,
                    "check_name": check_item.name,
                    "check_description": check_item.description,
                    "check_phase": check_item.phase,
                    "current_reliability": result.reliability,
                    "is_met_status": is_met_status,
                    "analysis_details": result.analysis_details,
                    "sources_summary": sources_summary
                }
                
                # Generate questions
                questions = llm.invoke(question_prompt.format(**question_input)).content
                print(f"\nNeed more information to evaluate this check. Please answer these questions:\n{questions}")
                
                # Get user input
                print("\nEnter your responses (type 'skip' to continue without additional info):")
                user_input = input("> ")
                
                if user_input.lower() == "skip":
                    print("Skipping additional information gathering.")
                    break
                
                # Add user input to additional info for next iteration
                additional_info += f"\n\n**Additional Information (Attempt {attempt+1}):**\n{user_input}"
                attempt += 1
            else:
                # Either reliability is good enough or we've reached max attempts
                break
                
        # Apply human review logic based on reliability after all attempts
        if result.reliability < 50:
            print(f"-> Check ID {check_item.id}: Final reliability {result.reliability:.1f}% < 50%. Flagging for human review.")
            result.needs_human_review = True
        else:
            print(f"-> Check ID {check_item.id}: Final reliability {result.reliability:.1f}% >= 50%. Looks OK.")
            result.needs_human_review = False
            
        return result
    except OutputParserException as ope:
        print(f"ERROR parsing LLM output for check ID {check_item.id}: {ope}")
        # Attempt to create a partial result indicating the failure
        return CheckResult(
            check_item=check_item, is_met=None, reliability=0.0, sources=[],
            analysis_details=f"Analysis failed: LLM output parsing error. Raw output might be: {ope.llm_output}",
            needs_human_review=True
        )
    except Exception as e:
        print(f"ERROR during RAG analysis for check ID {check_item.id}: {e}")
        traceback.print_exc()
        # Return a result indicating failure, flagged for review
        return CheckResult(
            check_item=check_item, is_met=None, reliability=0.0, sources=[],
            analysis_details=f"Analysis failed due to runtime error: {e}",
            needs_human_review=True
        )

def format_final_output(state: GraphState) -> GraphState:
    """Formats the results list into the final JSON structure and handles errors."""
    print(f"\n--- Node: format_final_output ---")

    if error_message := state.get("error_message"):
        print(f"Workflow finished with error: {error_message}")
        # Provide error in the final output
        state["final_results"] = [{"error": error_message}]
        return state

    analysis_results = state.get("analysis_results")

    if analysis_results is None:
         # This case might happen if decide_after_indexing skipped analysis
         # because there were no checks for the phase.
         if not state.get("checks_for_phase"):
             message = f"No checklist items found for the specified phase '{state['target_phase']}'. No analysis performed."
             print(message)
             state["final_results"] = [{"message": message}]
         else:
             # Should ideally not happen if map ran, but handle defensively
             print("Warning: Analysis results list is missing, but checks existed.")
             state["final_results"] = [{"warning": "Analysis results are unexpectedly missing."}]
         return state

    if not isinstance(analysis_results, list):
        print(f"Warning: 'analysis_results' is not a list (type: {type(analysis_results)}). Attempting to format anyway.")
        # Handle potential malformed state if map didn't return a list
        state["final_results"] = [{"error": "Internal state error: analysis_results is not a list."}]
        return state


    # Convert Pydantic models to dictionaries for JSON serialization
    output_list = []
    valid_results_count = 0
    human_review_count = 0

    for i, result in enumerate(analysis_results):
        if isinstance(result, CheckResult):
            output_list.append(result.dict(exclude_none=True)) # Exclude None fields for cleaner JSON
            valid_results_count += 1
            if result.needs_human_review:
                human_review_count += 1
        else:
            print(f"Warning: Item at index {i} in analysis_results is not a CheckResult object: {result}")
            output_list.append({"error": "Invalid result object received from analysis step.", "details": str(result)})


    state["final_results"] = output_list
    print(f"Formatted {valid_results_count} valid results.")

    # Log summary about human review items
    if valid_results_count > 0:
        if human_review_count > 0:
            print(f"SUMMARY: {human_review_count} out of {valid_results_count} analyzed checks require human review (Reliability < 50%).")
        else:
            print(f"SUMMARY: All {valid_results_count} analyzed checks met the reliability threshold (>= 50%).")
    elif not state.get("checks_for_phase"):
        # Message already handled above if checks_for_phase was empty
        pass
    else:
        print("SUMMARY: No valid results were generated.")


    return state

# --- Graph Definition ---

def decide_after_indexing(state: GraphState) -> str:
    """Determines the next step after document indexing."""
    print(f"\n--- Edge: decide_after_indexing ---")
    if state.get("error_message"):
        print("Decision: Error detected, routing to format_output.")
        return "format_output" # Go directly to format error output

    checks = state.get("checks_for_phase")
    if not checks: # Checks could be None or empty list
        print("Decision: No checks found for the phase. Routing to format_output.")
        return "format_output" # Skip analysis if no checks

    print(f"Decision: {len(checks)} checks found. Routing to analyze_checks_map.")
    return "analyze_checks_map" # Proceed to analysis via map

# --- Build the Graph ---
graph_builder = StateGraph(GraphState)

# Add nodes
graph_builder.add_node("load_filter_checklist", load_and_filter_checklist)
graph_builder.add_node("load_index_document", load_index_document)

# Define the node that performs the map operation
# It takes the state, extracts 'checks_for_phase', maps 'analyze_check_rag' over it,
# and the result (a list of CheckResult) is implicitly put into 'analysis_results'
# by the final lambda merging the dict back into the state.
analyze_check_node = RunnableLambda(analyze_check_rag, name="AnalyzeSingleCheck") # Give it a name for tracing

def analyze_checks_map_function(state: GraphState) -> GraphState:
    """Takes all checks and maps the analyze function over them with proper state config."""
    print(f"\n--- Node: analyze_checks_map function ---")
    checks = state.get("checks_for_phase", [])
    print(f"Processing {len(checks)} checks in map function...")
    
    # Create a config with the current state for each analysis
    config = {"configurable": state}
    
    # Map the analysis function over each check with proper configuration
    results = []
    for check in checks:
        result = analyze_check_node.invoke(check, config=config)
        results.append(result)
    
    # Update the state with results
    state["analysis_results"] = results
    return state

graph_builder.add_node("analyze_checks_map", analyze_checks_map_function)

graph_builder.add_node("format_output", format_final_output)

# Define edges
graph_builder.set_entry_point("load_filter_checklist")
graph_builder.add_edge("load_filter_checklist", "load_index_document")

# Conditional edge after indexing
graph_builder.add_conditional_edges(
    "load_index_document",
    decide_after_indexing,
    {
        "format_output": "format_output",      # If error or no checks
        "analyze_checks_map": "analyze_checks_map" # If checks exist and no error
    }
)

# After the map operation completes, go to format the output
graph_builder.add_edge("analyze_checks_map", "format_output")

# The final node leads to the end
graph_builder.add_edge("format_output", END)

# Compile the graph
try:
    app = graph_builder.compile()
    print("\n--- LangGraph Compiled Successfully ---")

    # Optional: Visualize the graph
    try:
        # Needs graphviz installed: pip install graphviz
        # app.get_graph().print_ascii()
        # Or save as image:
        # app.get_graph().draw_mermaid_png(output_file_path="graph.png")
        pass # Keep commented out unless graphviz is installed
    except Exception as viz_error:
        print(f"Could not visualize graph (is graphviz installed?): {viz_error}")

except Exception as compile_error:
    print(f"FATAL: LangGraph compilation failed: {compile_error}")
    traceback.print_exc()
    exit(1)


# --- Main Execution Block ---
if __name__ == "__main__":

    # --- Create Dummy Files for Demonstration ---
    checklist_file = "checklists/project_checklist_demo.xlsx"
    document_file = "input/quotation_demo.docx"

    # Create dummy checklist if it doesn't exist
    if not os.path.exists(checklist_file):
        print(f"\nCreating dummy checklist file: {checklist_file}")
        dummy_data = {
            'ID': [1, 1, 8, 8, 9, 9, 10, 11],
            'Name': ['Layout Validated', 'Layout Validated', 'Customer Approval', 'Customer Approval', 'Charger Location', 'Charger Location', 'Safety Certs', 'Network Ports'],
            'BranchID': [85, 85, 85, 85, 85, 85, 90, 95],
            'BranchName': ['LAYOUT', 'LAYOUT', 'LAYOUT', 'LAYOUT', 'LAYOUT', 'LAYOUT', 'SAFETY', 'IT'],
            'CHK_Description': [
                'Plant layout validated for clearances (1.5m+), obstacles, doors, and ceiling specifications (4m+).',
                'Plant layout validated for clearances (1.5m+), obstacles, doors, and ceiling specifications (4m+).',
                'Customer approval received for the final offer layout drawing.',
                'Customer approval received for the final offer layout drawing.',
                'Battery charger location defined and clearly marked (balooned) in the layout drawing.',
                'Battery charger location defined and clearly marked (balooned) in the layout drawing.',
                'Relevant safety certifications (e.g., CE marking) for major components are documented.',
                'Required network ports identified and locations specified in IT plan.'
            ],
            'Weight': [7, 10, 3, 5, 3, 10, 8, 5],
            'Phase': ['Apertura Commessa', 'Lancio', 'Apertura Commessa', 'Lancio', 'Apertura Commessa', 'Rilascio Tecnico', 'Apertura Commessa', 'Apertura Commessa']
        }
        try:
            pd.DataFrame(dummy_data).to_excel(checklist_file, index=False)
            print("Dummy checklist created.")
        except Exception as fe:
            print(f"Error creating dummy checklist: {fe}")

    # Create dummy document if it doesn't exist
    if not os.path.exists(document_file):
         print(f"Creating dummy document file: {document_file}")
         try:
             from docx import Document as DocxDocument
             doc = DocxDocument()
             doc.add_heading("Project Alpha - Quotation Details", level=1)
             doc.add_paragraph("Date: 2025-03-27")

             doc.add_heading("1. Layout Confirmation", level=2)
             doc.add_paragraph(
                 "The final layout drawing (Rev D, dated 2025-03-20) has been reviewed. "
                 "Main pathway clearances are confirmed at 1.6m. Ceiling height in the operational area is 4.1 meters. "
                 "All doorways meet standard industrial requirements. Known obstacles like pillars are marked."
             ) # Meets Check 1

             doc.add_heading("2. Approvals and Status", level=2)
             doc.add_paragraph(
                 "Following review meetings, the customer (Mr. Smith) provided verbal agreement on the layout Rev D during the call on March 22nd. "
                 "Formal written sign-off is expected by end of week."
             ) # Does NOT meet Check 8 (no formal approval received *yet*)

             doc.add_heading("3. Power Systems", level=2)
             doc.add_paragraph(
                 "The battery charging station area is designated near the maintenance bay. "
                 "Refer to drawing LAY-101 Rev D, balloon reference 'BC-01' for the exact placement."
             ) # Meets Check 9

             doc.add_heading("4. Compliance and IT", level=2)
             doc.add_paragraph(
                 "All major system components will be CE marked. Documentation packages containing certifications will be compiled and delivered during the commissioning phase."
             ) # Does NOT meet Check 10 (certs not documented *yet*)
             doc.add_paragraph(
                 "Network requirements are still under discussion with the customer's IT department. Port locations TBD."
             ) # Does NOT meet Check 11

             doc.save(document_file)
             print("Dummy document created.")
         except ImportError:
             print(f"Error creating dummy document: python-docx package not installed. Run 'pip install python-docx' to install it.")
         except Exception as de:
             print(f"Error creating dummy document: {de}")
             traceback.print_exc()


    # --- Define Inputs for the Graph Run ---
    # <<< --- MODIFY THESE INPUTS AS NEEDED --- >>>
    inputs = GraphState(
        checklist_path=checklist_file,
        document_path=document_file,
        target_phase="Apertura Commessa" # Analyze this specific phase
    )
    # <<< --- END OF MODIFIABLE INPUTS --- >>>


    print("\n" + "="*50)
    print("--- Starting Checklist RAG Workflow ---")
    print(f"Checklist: {inputs['checklist_path']}")
    print(f"Document: {inputs['document_path']}")
    print(f"Target Phase: {inputs['target_phase']}")
    print("="*50 + "\n")

    # --- Execute the Graph ---
    final_state = None
    try:
        # Configuration for the run, including recursion limit
        config = RunnableConfig(recursion_limit=25)

        # Invoke the graph
        # Use stream to see events (optional, good for debugging)
        # for event in app.stream(inputs, config=config):
        #     event_name = list(event.keys())[0]
        #     event_data = event[event_name]
        #     print(f"--- Event: {event_name} ---")
        #     # print(f"Data: {event_data}") # Can be verbose
        #     final_state = event_data # Keep track of the latest state

        # Or just invoke for the final result
        final_state = app.invoke(inputs, config=config)


        # --- Process Final Output ---
        print("\n" + "="*50)
        print("--- Workflow Completed ---")
        print("="*50)

        if final_state and final_state.get("final_results"):
            print("\nFinal Analysis Results (JSON):")
            # Pretty print the JSON output
            print(json.dumps(final_state["final_results"], indent=2))

            # Optionally save to a file
            output_filename = f"output/analysis_results_{inputs['target_phase'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            try:
                with open(output_filename, 'w') as f:
                    json.dump(final_state["final_results"], f, indent=2)
                print(f"\nResults also saved to: {output_filename}")
                
                # Generate a PDF report
                try:
                    from report_generator import ChecklistReportGenerator
                    print("\n--- Generating PDF Report ---")
                    generator = ChecklistReportGenerator(output_filename)
                    pdf_path = generator.generate_pdf()
                    print(f"PDF Report generated: {pdf_path}")
                except ImportError as ie:
                    print(f"\nCouldn't generate PDF report: {ie}")
                    print("Make sure 'matplotlib', 'reportlab', and 'report_generator.py' are available.")
                except Exception as pdf_error:
                    print(f"\nError generating PDF report: {pdf_error}")
                    traceback.print_exc()
                
            except Exception as write_error:
                print(f"\nError saving results to file: {write_error}")

        elif final_state and final_state.get("error_message"):
             print(f"\nWorkflow finished with an error state: {final_state['error_message']}")
             # Optionally print the partial state for debugging
             # print("\nFinal State (Error):")
             # print(final_state)
        else:
            print("\n--- Workflow finished, but no final results found in the expected format. ---")
            print("Final State:")
            print(final_state)

    except Exception as e:
        print(f"\n--- FATAL: Workflow Execution Failed ---")
        print(f"Error Type: {type(e).__name__}")
        print(f"Error Details: {e}")
        traceback.print_exc()
        # Print state if available
        if final_state:
            print("\n--- State at time of error ---")
            print(final_state)

    print("\n--- End of script ---")