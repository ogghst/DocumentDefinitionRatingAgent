import os
import pandas as pd
import traceback
from docx import Document as DocxDocument

def create_demo_files():
    """Create demo files for the checklist RAG application if they don't exist."""
    checklist_file = "checklists/project_checklist_demo.xlsx"
    document_file = "input/quotation_demo.docx"
    
    # Create directories if they don't exist
    os.makedirs(os.path.dirname(checklist_file), exist_ok=True)
    os.makedirs(os.path.dirname(document_file), exist_ok=True)
    
    # Create dummy checklist if it doesn't exist
    if not os.path.exists(checklist_file):
        print(f"\nCreating dummy checklist file: {checklist_file}")
        dummy_data = {
            'ID': [1, 1, 8, 8, 9, 9, 10, 11],
            'Name': ['Layout Validated', 'Layout Validated', 'Customer Approval', 'Customer Approval', 'Charger Location', 'Charger Location', 'Safety Certs', 'Network Ports'],
            'BranchID': [85, 85, 85, 85, 85, 85, 90, 95],
            'BranchName': ['LAYOUT', 'LAYOUT', 'LAYOUT', 'LAYOUT', 'LAYOUT', 'LAYOUT', 'SAFETY', 'IT'],
            'CHK_Description': [
                'Plant layout validated for clearances (1.5m+), obstacles, doors, and ceiling specifications (4m+).',
                'Plant layout validated for clearances (1.5m+), obstacles, doors, and ceiling specifications (4m+).',
                'Customer approval received for the final offer layout drawing.',
                'Customer approval received for the final offer layout drawing.',
                'Battery charger location defined and clearly marked (balooned) in the layout drawing.',
                'Battery charger location defined and clearly marked (balooned) in the layout drawing.',
                'Relevant safety certifications (e.g., CE marking) for major components are documented.',
                'Required network ports identified and locations specified in IT plan.'
            ],
            'Weight': [7, 10, 3, 5, 3, 10, 8, 5],
            'Phase': ['Apertura Commessa', 'Lancio', 'Apertura Commessa', 'Lancio', 'Apertura Commessa', 'Rilascio Tecnico', 'Apertura Commessa', 'Apertura Commessa']
        }
        try:
            pd.DataFrame(dummy_data).to_excel(checklist_file, index=False)
            print("Dummy checklist created.")
        except Exception as fe:
            print(f"Error creating dummy checklist: {fe}")

    # Create dummy document if it doesn't exist
    if not os.path.exists(document_file):
        print(f"Creating dummy document file: {document_file}")
        try:
            doc = DocxDocument()
            doc.add_heading("Project Alpha - Quotation Details", level=1)
            doc.add_paragraph("Date: 2025-03-27")

            doc.add_heading("1. Layout Confirmation", level=2)
            doc.add_paragraph(
                "The final layout drawing (Rev D, dated 2025-03-20) has been reviewed. "
                "Main pathway clearances are confirmed at 1.6m. Ceiling height in the operational area is 4.1 meters. "
                "All doorways meet standard industrial requirements. Known obstacles like pillars are marked."
            )

            doc.add_heading("2. Approvals and Status", level=2)
            doc.add_paragraph(
                "Following review meetings, the customer (Mr. Smith) provided verbal agreement on the layout Rev D during the call on March 22nd. "
                "Formal written sign-off is expected by end of week."
            )

            doc.add_heading("3. Power Systems", level=2)
            doc.add_paragraph(
                "The battery charging station area is designated near the maintenance bay. "
                "Refer to drawing LAY-101 Rev D, balloon reference 'BC-01' for the exact placement."
            )

            doc.add_heading("4. Compliance and IT", level=2)
            doc.add_paragraph(
                "All major system components will be CE marked. Documentation packages containing certifications will be compiled and delivered during the commissioning phase."
            )
            doc.add_paragraph(
                "Network requirements are still under discussion with the customer's IT department. Port locations TBD."
            )

            doc.save(document_file)
            print("Dummy document created.")
        except Exception as de:
            print(f"Error creating dummy document: {de}")
            traceback.print_exc() 